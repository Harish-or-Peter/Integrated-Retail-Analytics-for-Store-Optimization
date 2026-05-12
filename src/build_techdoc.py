"""Convert docs/TECHNICAL_DOCUMENT.md into a Google-Docs-importable .docx file.

Run with: python src/build_techdoc.py
Output:   docs/TECHNICAL_DOCUMENT.docx
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path('docs/TECHNICAL_DOCUMENT.md')
OUT = Path('docs/TECHNICAL_DOCUMENT.docx')

doc = Document()

# Set base font.
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_runs(paragraph, text):
    """Add inline runs to a paragraph, parsing **bold**, *italic*, and `code`."""
    # Tokenize: split on the three styles in priority order.
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    for token in pattern.split(text):
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2]); run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1]); run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'; run.font.size = Pt(10)
        else:
            paragraph.add_run(token)


def add_table_from_md(rows):
    """rows is a list of '| a | b | c |' lines (header + alignment + body)."""
    # Strip outer pipes and split.
    parsed = []
    for line in rows:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        parsed.append(cells)
    header = parsed[0]
    body = parsed[2:]  # skip alignment row
    ncols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=ncols)
    table.style = 'Light Grid Accent 1'
    # Header row.
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        add_runs(p, cell_text)
        for run in p.runs:
            run.bold = True
    # Body.
    for i, row in enumerate(body):
        for j, cell_text in enumerate(row[:ncols]):
            cell = table.rows[i + 1].cells[j]
            cell.text = ''
            add_runs(cell.paragraphs[0], cell_text)


# Parse the markdown line by line.
lines = SRC.read_text(encoding='utf-8').splitlines()
i = 0
while i < len(lines):
    line = lines[i].rstrip()

    # Skip Markdown quote / divider lines.
    if line.startswith('---'):
        doc.add_paragraph().add_run().add_break()
        i += 1
        continue
    if line.startswith('> '):
        p = doc.add_paragraph(style='Intense Quote')
        add_runs(p, line[2:])
        i += 1
        continue

    # Headings.
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        text  = line[level:].strip()
        if level == 1:
            h = doc.add_heading(text, level=0)
        else:
            h = doc.add_heading(text, level=min(level - 1, 4))
        i += 1
        continue

    # Tables (start of one).
    if line.startswith('|') and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i + 1].strip()):
        block = [line]
        i += 1
        while i < len(lines) and lines[i].strip().startswith('|'):
            block.append(lines[i].strip())
            i += 1
        add_table_from_md(block)
        continue

    # Bullets.
    if line.lstrip().startswith('- '):
        text = line.lstrip()[2:]
        p = doc.add_paragraph(style='List Bullet')
        add_runs(p, text)
        i += 1
        continue

    # Numbered list (1. 2. ...).
    if re.match(r'^\d+\.\s', line.lstrip()):
        text = re.sub(r'^\d+\.\s', '', line.lstrip())
        p = doc.add_paragraph(style='List Number')
        add_runs(p, text)
        i += 1
        continue

    # Blank line.
    if not line.strip():
        doc.add_paragraph()
        i += 1
        continue

    # Default paragraph.
    p = doc.add_paragraph()
    add_runs(p, line)
    i += 1

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f'Wrote {OUT}  ({OUT.stat().st_size:,} bytes)')
